# from config import CAM_REPORT_TITLE, CAM_AUTHOR, OUTPUTS_DIR
"""
cam_generator.py  —  IntelliCredit v2.0
INVESTMENT ASSESSMENT REPORT  —  PDF + DOCX

TARGET: 5-6 pages maximum
- Page 1: Cover (white bg, full entity + loan details, IntelliCredit branding)
- Page 2: Executive Summary + Financial Analysis + Five Cs
- Page 3: AI Risk Drivers + Reasoning + SWOT
- Page 4: External Research + Triangulation + Early Warnings
- Page 5: Conditions + Recommendation + Sign-off

COLOUR SYSTEM
─────────────
  Navy  #0d1f5c  — structural chrome, headers, borders
  Gold  #c9970a  — cover accent only
  Green #1a6b2a  — APPROVE
  Amber #b85c00  — CONDITIONAL + warnings
  Red   #b71c1c  — REJECT + adverse
  White #ffffff  — text on dark
  LBlue #e8edf8  — label column tint
  Tint  #f5f7fd  — alternating rows
  MGray #c6cad8  — grid lines
  DGray #4c5068  — footer / captions
"""

from config import CAM_REPORT_TITLE, CAM_AUTHOR, OUTPUTS_DIR
from src.schemas import CreditAppraisalResult
from datetime import datetime
import sys
import math
from pathlib import Path

ROOT_DIR = Path(__file__).parent.parent
sys.path.insert(0, str(ROOT_DIR))

# ── Palette ───────────────────────────────────────────────────────────────────
NAV = "#0d1f5c"
NAV2 = "#1a3080"
GOLD = "#c9970a"
LBLU = "#e8edf8"
TINT = "#f5f7fd"
MGRY = "#c6cad8"
DGRY = "#4c5068"
GRN = "#1a6b2a"
AMB = "#b85c00"
RED = "#b71c1c"
WHT = "#ffffff"
GRN_L = "#e8f5ea"
AMB_L = "#fff4e5"
RED_L = "#ffeaea"

APP_NAME = "IntelliCredit"          # ← service name, NOT bank name


class CAMGenerator:
    """Generates INVESTMENT ASSESSMENT REPORT in PDF and DOCX — 5-6 pages."""

    # ── public ────────────────────────────────────────────────────────────────

    def generate_both(self, result: CreditAppraisalResult) -> dict:
        safe = result.company_name.replace(" ", "_").replace("/", "-")
        ts = datetime.now().strftime("%Y%m%d_%H%M")
        pdf = str(OUTPUTS_DIR / f"CAM_{safe}_{ts}.pdf")
        docx = str(OUTPUTS_DIR / f"CAM_{safe}_{ts}.docx")
        print("Generating PDF CAM...")
        self.generate_pdf(result, pdf)
        print("Generating DOCX CAM...")
        self.generate_docx(result, docx)
        print(f"Reports saved: {OUTPUTS_DIR}")
        return {"pdf": pdf, "docx": docx}

    # ── shared helpers ────────────────────────────────────────────────────────

    def _pred_strings(self, pred):
        if not pred:
            return "PENDING", "UNKNOWN", WHT, WHT, "PENDING"
        ds = str(pred.decision).replace("DecisionType.", "").upper()
        cs = str(pred.risk_category).replace("RiskCategory.", "").upper()
        if "APPROVE" in ds:
            bg, lt, label = GRN, GRN_L, "APPROVED"
        elif "REJECT" in ds:
            bg, lt, label = RED, RED_L, "REJECTED"
        else:
            bg, lt, label = AMB, AMB_L, "CONDITIONAL APPROVAL"
        return ds, cs, bg, lt, label

    def _derived_rows(self, result):
        d = result.derived_financials
        if not d:
            return []
        rows = []
        if d.debt_equity_ratio is not None:
            rows.append(("D/E Ratio",           f"{d.debt_equity_ratio:.2f}x"))
        if d.net_worth_inr is not None:
            rows.append(
                ("Net Worth",            f"Rs. {d.net_worth_inr:,.0f}"))
        if d.dscr is not None:
            rows.append(("DSCR",                 f"{d.dscr:.2f}x"))
        if d.net_profit_margin is not None:
            rows.append(
                ("Net Profit Margin",    f"{d.net_profit_margin:.1f}%"))
        if getattr(d, "avg_monthly_balance_inr", None) is not None:
            rows.append(
                ("Avg Monthly Bal.",   f"Rs. {d.avg_monthly_balance_inr:,.0f}"))
        if d.data_completeness_pct is not None:
            rows.append(
                ("Data Completeness",  f"{d.data_completeness_pct:.0f}%"))
        return rows

    # ── risk gauge ────────────────────────────────────────────────────────────

    def _risk_gauge(self, score: float, W=160, H=100):
        from reportlab.graphics.shapes import Drawing, Wedge, Circle, String, Line
        from reportlab.lib import colors as rc
        d = Drawing(W, H)
        cx = W / 2
        cy = H * 0.10
        ro = H * 0.85
        ri = H * 0.50
        rn = H * 0.72
        for sa, ea, col in [(180, 120, GRN), (120, 60, "#e8a020"), (60, 0, RED)]:
            d.add(Wedge(cx, cy, ro, ea, sa, fillColor=rc.HexColor(
                col),   strokeColor=rc.white, strokeWidth=2))
            d.add(Wedge(cx, cy, ri, ea, sa, fillColor=rc.white,
                  strokeColor=rc.white, strokeWidth=1))
        ang = math.radians(180 - score*180)
        nx = cx + rn*math.cos(ang)
        ny = cy + rn*math.sin(ang)
        d.add(Line(cx+1, cy-1, nx+1, ny-1,
              strokeColor=rc.HexColor("#aaaacc"), strokeWidth=2))
        d.add(Line(cx, cy, nx, ny,         strokeColor=rc.HexColor(
            NAV),        strokeWidth=3))
        d.add(Circle(cx, cy, 6, fillColor=rc.HexColor(
            NAV), strokeColor=rc.white, strokeWidth=1.5))
        sc = GRN if score < 0.33 else ("#e8a020" if score < 0.66 else RED)
        d.add(String(cx, cy+ri*0.15, f"{score:.3f}", fontSize=13, fontName="Helvetica-Bold",
                     fillColor=rc.HexColor(sc), textAnchor="middle"))
        for lbl, ang_d, col in [("LOW", 150, GRN), ("MED", 90, "#e8a020"), ("HIGH", 30, RED)]:
            mr = (ro+ri)/2
            d.add(String(cx+mr*math.cos(math.radians(ang_d)),
                         cy+mr*math.sin(math.radians(ang_d))-3, lbl,
                         fontSize=7, fontName="Helvetica-Bold",
                         fillColor=rc.white, textAnchor="middle"))
        return d

    # ══════════════════════════════════════════════════════════════════════════
    # PDF
    # ══════════════════════════════════════════════════════════════════════════

    def generate_pdf(self, result: CreditAppraisalResult, output_path: str) -> str:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
        from reportlab.platypus import (
            BaseDocTemplate, Frame, PageTemplate,
            Paragraph, Spacer, Table, TableStyle,
            HRFlowable, PageBreak, NextPageTemplate,
        )
        from reportlab.graphics.shapes import Drawing, Rect, String, Line

        PW, PH = A4
        LM = RM = 1.7*cm
        TM = 2.2*cm
        BM = 1.8*cm
        TW = PW - LM - RM

        C = colors.HexColor
        cNAV = C(NAV)
        cNAV2 = C(NAV2)
        cGOLD = C(GOLD)
        cLBL = C(LBLU)
        cTNT = C(TINT)
        cMGR = C(MGRY)
        cDGR = C(DGRY)
        cGRN = C(GRN)
        cAMB = C(AMB)
        cRED = C(RED)
        cWHT = colors.white

        SS = getSampleStyleSheet()
        _id = [0]

        def PS(base="Normal", **kw):
            _id[0] += 1
            base_s = kw.pop("parent", SS.get(base, SS["Normal"]))
            return ParagraphStyle(f"_s{_id[0]}", parent=base_s, **kw)

        pN = PS(fontSize=8.5, fontName="Helvetica",
                leading=12.5, textColor=colors.black)
        pB = PS(fontSize=8.5, fontName="Helvetica-Bold",
                leading=12.5, textColor=colors.black)
        pS = PS(fontSize=7,   fontName="Helvetica",
                leading=10.5, textColor=cDGR)
        pNC = PS(fontSize=8.5, fontName="Helvetica",      leading=12.5,
                 alignment=TA_CENTER, textColor=colors.black)
        pBC = PS(fontSize=8.5, fontName="Helvetica-Bold", leading=12.5,
                 alignment=TA_CENTER, textColor=colors.black)
        pWC = PS(fontSize=8.5, fontName="Helvetica-Bold",
                 leading=12.5, alignment=TA_CENTER, textColor=cWHT)
        pWL = PS(fontSize=8.5, fontName="Helvetica-Bold",
                 leading=12.5, textColor=cWHT)
        pJ = PS(fontSize=8.5, fontName="Helvetica",      leading=13,
                alignment=TA_JUSTIFY, textColor=colors.black)

        pred = result.risk_prediction
        ds, cs, dec_bg, dec_lt, dec_label = self._pred_strings(pred)
        now = datetime.now()
        ref = f"CAM/{now.year}/{result.company_name[:4].upper()}/{now.strftime('%m%d%H%M')}"

        # ── page callbacks ────────────────────────────────────────────────────

        def _cover_bg(canv, doc):
            canv.saveState()
            # white background — full page
            canv.setFillColor(colors.white)
            canv.rect(0, 0, PW, PH, fill=1, stroke=0)
            # navy top bar (2.8 cm)
            canv.setFillColor(cNAV)
            canv.rect(0, PH-2.8*cm, PW, 2.8*cm, fill=1, stroke=0)
            # APP_NAME in bar
            canv.setFont("Helvetica-Bold", 15)
            canv.setFillColor(cWHT)
            canv.drawCentredString(PW/2, PH-1.6*cm, APP_NAME.upper())
            canv.setFont("Helvetica", 8)
            canv.setFillColor(C("#b8c8f0"))
            canv.drawCentredString(
                PW/2, PH-2.3*cm, "AI-Powered Credit Intelligence Engine")
            # gold rule
            canv.setStrokeColor(cGOLD)
            canv.setLineWidth(2)
            canv.line(LM, PH-2.95*cm, PW-RM, PH-2.95*cm)
            # navy bottom strip
            canv.setFillColor(cNAV)
            canv.rect(0, 0, PW, 1.1*cm, fill=1, stroke=0)
            canv.setFont("Helvetica", 6.5)
            canv.setFillColor(C("#b8c8f0"))
            canv.drawCentredString(
                PW/2, 0.38*cm, "STRICTLY CONFIDENTIAL  —  FOR INTERNAL USE ONLY")
            # watermark
            canv.saveState()
            canv.translate(PW/2, PH/2)
            canv.rotate(35)
            canv.setFont("Helvetica-Bold", 54)
            canv.setFillColor(colors.Color(0.85, 0.87, 0.93, alpha=0.08))
            canv.drawCentredString(0, 0, "CONFIDENTIAL")
            canv.restoreState()
            canv.restoreState()

        def _body_bg(canv, doc):
            canv.saveState()
            canv.setFillColor(cNAV)
            canv.rect(0, PH-1.2*cm, PW, 1.2*cm, fill=1, stroke=0)
            canv.setFont("Helvetica-Bold", 7.5)
            canv.setFillColor(cWHT)
            canv.drawString(LM, PH-0.75*cm, APP_NAME.upper())
            canv.setFont("Helvetica", 7)
            canv.setFillColor(C("#b8c8f0"))
            canv.drawRightString(
                PW-RM, PH-0.75*cm, f"Ref: {ref}   |   Page {doc.page}")
            canv.setStrokeColor(cGOLD)
            canv.setLineWidth(1.2)
            canv.line(0, PH-1.3*cm, PW, PH-1.3*cm)
            canv.setStrokeColor(cMGR)
            canv.setLineWidth(0.5)
            canv.line(LM, 1.3*cm, PW-RM, 1.3*cm)
            canv.setFont("Helvetica", 6.5)
            canv.setFillColor(cDGR)
            canv.drawString(
                LM, 0.65*cm, "STRICTLY CONFIDENTIAL  —  FOR INTERNAL USE ONLY")
            canv.drawRightString(
                PW-RM, 0.65*cm, f"{now.strftime('%d %b %Y')}   |   {CAM_AUTHOR}")
            canv.restoreState()

        # ── frames & templates ────────────────────────────────────────────────
        cover_frame = Frame(LM, 1.2*cm, TW, PH-3.1*cm-1.2 *
                            cm, id="cover", showBoundary=0)
        body_frame = Frame(LM, BM,     TW, PH-TM-BM,
                           id="body",  showBoundary=0)
        pdf_doc = BaseDocTemplate(output_path, pagesize=A4,
                                  leftMargin=LM, rightMargin=RM, topMargin=TM, bottomMargin=BM)
        pdf_doc.addPageTemplates([
            PageTemplate(id="Cover", frames=[cover_frame], onPage=_cover_bg),
            PageTemplate(id="Body",  frames=[body_frame],  onPage=_body_bg),
        ])

        # ── helpers ───────────────────────────────────────────────────────────

        def info_tbl(rows, lw=4.5*cm, fs=8.5, pad=4):
            rw = TW - lw
            data = [[Paragraph(str(l), PS(fontSize=fs, fontName="Helvetica-Bold",
                                          leading=fs+3.5, textColor=colors.black)),
                     Paragraph(str(v), PS(fontSize=fs, fontName="Helvetica",
                                          leading=fs+3.5, textColor=colors.black))]
                    for l, v in rows]
            t = Table(data, colWidths=[lw, rw])
            t.setStyle(TableStyle([
                ("BACKGROUND",    (0, 0), (0, -1), cLBL),
                ("ROWBACKGROUNDS", (1, 0), (1, -1), [cWHT, cTNT]),
                ("BOX",           (0, 0), (-1, -1), 0.6, cNAV2),
                ("INNERGRID",     (0, 0), (-1, -1), 0.25, cMGR),
                ("LEFTPADDING",   (0, 0), (-1, -1), pad+1),
                ("RIGHTPADDING",  (0, 0), (-1, -1), pad),
                ("TOPPADDING",    (0, 0), (-1, -1), pad),
                ("BOTTOMPADDING", (0, 0), (-1, -1), pad),
                ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
            ]))
            return t

        def shdr(title):
            drw = Drawing(TW, 16)
            drw.add(Rect(0, 0, 3.5, 16, fillColor=cNAV, strokeColor=None))
            drw.add(Line(0, 0, TW, 0, strokeColor=cMGR, strokeWidth=0.5))
            drw.add(String(9, 3, title, fontSize=9.5, fontName="Helvetica-Bold",
                           fillColor=C(NAV2)))
            return drw

        def col_tbl(rows, widths, hdr=NAV):
            t = Table(rows, colWidths=widths, repeatRows=1)
            t.setStyle(TableStyle([
                ("BACKGROUND",    (0, 0), (-1, 0),  C(hdr)),
                ("TEXTCOLOR",     (0, 0), (-1, 0),  cWHT),
                ("FONTNAME",      (0, 0), (-1, 0),  "Helvetica-Bold"),
                ("FONTSIZE",      (0, 0), (-1, -1),  8),
                ("LEADING",       (0, 0), (-1, -1),  11.5),
                ("BOX",           (0, 0), (-1, -1),  0.6, cNAV2),
                ("INNERGRID",     (0, 0), (-1, -1),  0.25, cMGR),
                ("LEFTPADDING",   (0, 0), (-1, -1),  5),
                ("RIGHTPADDING",  (0, 0), (-1, -1),  5),
                ("TOPPADDING",    (0, 0), (-1, -1),  4),
                ("BOTTOMPADDING", (0, 0), (-1, -1),  4),
                ("VALIGN",        (0, 0), (-1, -1),  "MIDDLE"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1),  [cWHT, cTNT]),
            ]))
            return t

        def gap(h): return Spacer(1, h*cm)

        # ══════════════════════════════════════════════════════════════════════
        # PAGE 1 — COVER
        # ══════════════════════════════════════════════════════════════════════
        story = [NextPageTemplate("Cover")]
        story.append(gap(0.5))

        story.append(Paragraph("INVESTMENT ASSESSMENT REPORT",
                               PS(fontSize=19, fontName="Helvetica-Bold", alignment=TA_CENTER,
                                  textColor=C(NAV), spaceAfter=4)))
        story.append(Paragraph("Internal Document  —  For Investment Committee",
                               PS(fontSize=8, fontName="Helvetica", alignment=TA_CENTER,
                                  textColor=cDGR, spaceAfter=0)))
        story.append(gap(0.2))
        story.append(HRFlowable(width="70%", thickness=1.6,
                     color=cGOLD, hAlign="CENTER"))
        story.append(gap(0.35))

        # ── Two-column cover table: Entity | Loan ────────────────────────────
        _cin = getattr(result, "cin", None) or "N/A"
        _pan = (result.itr_data.pan if result.itr_data and result.itr_data.pan
                else getattr(result, "pan", "N/A"))
        _gstin = (result.gst_data.gstin if result.gst_data and result.gst_data.gstin
                  else "N/A")
        _sector = getattr(result, "sector", None) or "—"
        _ltype = getattr(result, "loan_type", None) or "—"
        _lamount = getattr(result, "loan_amount_cr", None)
        _tenure = getattr(result, "loan_tenure_months", None)

        entity_rows = [
            ("Company / Borrower", result.company_name),
            ("CIN",    _cin),
            ("PAN",    _pan),
            ("GSTIN",  _gstin),
            ("Sector", _sector),
        ]
        loan_rows = [
            ("Loan Type",    _ltype),
            ("Amount Req.",
             f"Rs. {_lamount:.2f} Cr" if _lamount else "As per application"),
            ("Tenure",
             f"{_tenure} months" if _tenure else "As per sanction"),
            ("AI Decision",  dec_label),
            ("Risk Score",   f"{pred.risk_score:.3f}" if pred else "N/A"),
        ]

        def mini_tbl(rows, lw=3.0*cm):
            rw = TW/2 - lw - 0.3*cm
            data = [[Paragraph(str(l), PS(fontSize=8, fontName="Helvetica-Bold",
                                          leading=11, textColor=colors.black)),
                     Paragraph(str(v), PS(fontSize=8, fontName="Helvetica",
                                          leading=11, textColor=colors.black))]
                    for l, v in rows]
            t = Table(data, colWidths=[lw, rw])
            t.setStyle(TableStyle([
                ("BACKGROUND",    (0, 0), (0, -1), cLBL),
                ("ROWBACKGROUNDS", (1, 0), (1, -1), [cWHT, cTNT]),
                ("BOX",           (0, 0), (-1, -1), 0.5, cNAV2),
                ("INNERGRID",     (0, 0), (-1, -1), 0.2, cMGR),
                ("LEFTPADDING",   (0, 0), (-1, -1), 5),
                ("RIGHTPADDING",  (0, 0), (-1, -1), 5),
                ("TOPPADDING",    (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
            ]))
            return t

        cover_cols = Table(
            [[mini_tbl(entity_rows), mini_tbl(loan_rows)]],
            colWidths=[TW/2, TW/2]
        )
        cover_cols.setStyle(TableStyle([
            ("LEFTPADDING",   (0, 0), (-1, -1), 0),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 3),
            ("TOPPADDING",    (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ]))
        story.append(cover_cols)
        story.append(gap(0.3))

        # ── Decision banner ───────────────────────────────────────────────────
        banner_data = [[
            Paragraph(f"<b>{dec_label}</b>",
                      PS(fontSize=12, fontName="Helvetica-Bold",
                         alignment=TA_CENTER, textColor=cWHT))
        ]]
        banner = Table(banner_data, colWidths=[TW])
        banner.setStyle(TableStyle([
            ("BACKGROUND",    (0, 0), (-1, -1), C(dec_bg)),
            ("TOPPADDING",    (0, 0), (-1, -1), 10),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
        ]))
        story.append(banner)
        story.append(gap(0.3))

        # ── Gauge + key metrics row ───────────────────────────────────────────
        if pred:
            gauge = self._risk_gauge(pred.risk_score, W=130, H=84)
            metrics_rows = [
                ("Risk Score",    f"{pred.risk_score:.3f} / 1.000"),
                ("Risk Category", cs),
                ("Loan Limit",    f"Rs. {pred.loan_limit_inr:,.0f}"),
                ("Interest Rate", f"{pred.interest_rate:.2f}% p.a."),
                ("Reference",     ref),
                ("Report Date",   now.strftime("%d %B %Y")),
            ]
            metrics_tbl = mini_tbl(metrics_rows, lw=3.2*cm)
            gauge_wrap = Table([[gauge]], colWidths=[4.5*cm])
            gauge_wrap.setStyle(TableStyle([
                ("ALIGN",  (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("BOX",    (0, 0), (-1, -1), 0.5, cMGR),
                ("BACKGROUND", (0, 0), (-1, -1), cTNT),
                ("TOPPADDING",    (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]))
            row_wrap = Table([[gauge_wrap, metrics_tbl]],
                             colWidths=[4.8*cm, TW-4.8*cm])
            row_wrap.setStyle(TableStyle([
                ("LEFTPADDING",   (0, 0), (-1, -1), 0),
                ("RIGHTPADDING",  (0, 0), (-1, -1), 4),
                ("TOPPADDING",    (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                ("VALIGN",        (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(row_wrap)
            story.append(gap(0.3))

        # ── Signature block ───────────────────────────────────────────────────
        cw3 = TW / 3
        sig_data = [
            [Paragraph("Prepared by", pWC), Paragraph(
                "Reviewed by", pWC), Paragraph("Approved by", pWC)],
            [Paragraph("Credit Analyst", pNC), Paragraph(
                "Credit Manager", pNC), Paragraph("Chief Credit Officer", pNC)],
            [Paragraph("\n\n___________________", pNC)]*3,
            [Paragraph("Date: _______________", pS)]*3,
        ]
        sig_tbl = Table(sig_data, colWidths=[cw3, cw3, cw3])
        sig_tbl.setStyle(TableStyle([
            ("BOX",           (0, 0), (-1, -1), 0.7, cNAV),
            ("INNERGRID",     (0, 0), (-1, -1), 0.3, cMGR),
            ("BACKGROUND",    (0, 0), (-1, 1),  cNAV),
            ("BACKGROUND",    (0, 2), (-1, -1), cTNT),
            ("TOPPADDING",    (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING",   (0, 0), (-1, -1), 4),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 4),
            ("ALIGN",         (0, 0), (-1, -1), "CENTER"),
        ]))
        story.append(sig_tbl)

        # ── Body pages ────────────────────────────────────────────────────────
        story.append(NextPageTemplate("Body"))
        story.append(PageBreak())

        # ══════════════════════════════════════════════════════════════════════
        # PAGE 2 — FINANCIAL ANALYSIS + FIVE Cs
        # ══════════════════════════════════════════════════════════════════════
        story.append(
            shdr("SECTION 1   —   EXECUTIVE SUMMARY & FINANCIAL ANALYSIS"))
        story.append(gap(0.15))

        # Decisive factor
        if pred and getattr(pred, "decisive_factor", ""):
            df_t = Table([[Paragraph(
                f"<b>⚡ DECISIVE FACTOR:</b>  {pred.decisive_factor}", pJ)]],
                colWidths=[TW])
            df_t.setStyle(TableStyle([
                ("BOX",           (0, 0), (-1, -1), 0.8, cNAV2),
                ("BACKGROUND",    (0, 0), (-1, -1), cLBL),
                ("LEFTPADDING",   (0, 0), (-1, -1), 10),
                ("RIGHTPADDING",  (0, 0), (-1, -1), 10),
                ("TOPPADDING",    (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]))
            story.append(df_t)
            story.append(gap(0.15))

        # Two columns: Financial Ratios | GST + Bank
        col_left_rows = self._derived_rows(result)
        gst = result.gst_data
        bk = result.bank_data
        d = result.derived_financials
        col_right_rows = []
        if gst:
            col_right_rows.append(("GSTIN",         gst.gstin or "N/A"))
            col_right_rows.append(
                ("GST Turnover",   f"Rs.{gst.turnover:,.0f}" if gst.turnover else "N/A"))
            col_right_rows.append(
                ("ITC Claimed",    f"Rs.{gst.itc_claimed:,.0f}" if gst.itc_claimed else "N/A"))
        if bk:
            avg = getattr(bk, "average_monthly_balance", 0) or (
                getattr(d, "avg_monthly_balance_inr", 0) if d else 0)
            col_right_rows.append(
                ("Avg Monthly Bal.", f"Rs.{avg:,.0f}" if avg else "N/A"))
            col_right_rows.append(("EMI Bounces",    str(
                getattr(bk, "emi_bounce_count", 0) or 0)))
        rec = result.gst_reconciliation
        if rec:
            flag = "⚠️ RISK FLAG" if rec.risk_flag else "✅ Passed"
            col_right_rows.append(
                ("GST Reconciliation", f"{flag} — {rec.variance_pct:.1f}%"))

        while len(col_left_rows) < len(col_right_rows):
            col_left_rows.append(("", ""))
        while len(col_right_rows) < len(col_left_rows):
            col_right_rows.append(("", ""))

        if col_left_rows or col_right_rows:
            lh = mini_tbl(col_left_rows or [
                          ("—", "No financial data")], lw=2.8*cm)
            rh = mini_tbl(col_right_rows or [
                          ("—", "No GST/bank data")],  lw=2.8*cm)
            # Add BEFORE the two_col Table:
            story.append(Paragraph(
                "Financial Ratios &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;"
                "&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;"
                "&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;"
                "GST & Banking",
                PS(fontSize=8, fontName="Helvetica-Bold", textColor=C(NAV2),
                   spaceBefore=4, spaceAfter=4)
            ))
            two_col = Table([[lh, rh]], colWidths=[TW/2, TW/2])
            two_col.setStyle(TableStyle([
                ("LEFTPADDING", (0, 0), (-1, -1),
                 0), ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("TOPPADDING", (0, 0), (-1, -1),
                 0), ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(two_col)
        story.append(gap(0.25))

        # Five Cs table
        story.append(shdr("SECTION 2   —   FIVE Cs CREDIT ANALYSIS"))
        story.append(gap(0.12))
        if result.five_cs:
            obj5 = result.five_cs
            W0, W1, W2, W3 = 2.5*cm, 1.7*cm, 2.3*cm, TW-6.8*cm
            rows5 = [[Paragraph("Parameter", pWL), Paragraph("Score", pWC),
                      Paragraph("Rating", pWC), Paragraph("Assessment", pWL)]]
            for lbl, o in [("Character", obj5.character), ("Capacity", obj5.capacity),
                           ("Capital", obj5.capital), ("Collateral", obj5.collateral),
                           ("Conditions", obj5.conditions)]:
                sc = o.score
                rat, rc2 = (("Excellent", GRN) if sc >= 8.5 else ("Good", GRN) if sc >= 7 else
                            ("Adequate", AMB) if sc >= 5.5 else ("Weak", RED))
                rows5.append([Paragraph(lbl, pB), Paragraph(f"{sc}/10", pBC),
                              Paragraph(
                                  f'<font color="{rc2}"><b>{rat}</b></font>', pNC),
                              Paragraph(o.summary, pN)])
            rows5.append([Paragraph("<b>OVERALL</b>", pB),
                          Paragraph(f"<b>{obj5.overall_score}/10</b>", pBC),
                          Paragraph("<b>Wtd. Avg.</b>", pBC),
                          Paragraph("Combined weighted Five Cs score.", pN)])
            t5 = Table(rows5, colWidths=[W0, W1, W2, W3])
            t5.setStyle(TableStyle([
                ("BACKGROUND",    (0, 0), (-1, 0),  cNAV),
                ("TEXTCOLOR",     (0, 0), (-1, 0),  cWHT),
                ("BACKGROUND",    (0, -1), (-1, -1), cLBL),
                ("FONTNAME",      (0, -1), (-1, -1), "Helvetica-Bold"),
                ("FONTSIZE",      (0, 0), (-1, -1),
                 8), ("LEADING", (0, 0), (-1, -1), 11.5),
                ("BOX",           (0, 0), (-1, -1),  0.6, cNAV2),
                ("INNERGRID",     (0, 0), (-1, -1),  0.25, cMGR),
                ("LEFTPADDING",   (0, 0), (-1, -1),
                 5), ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING",    (0, 0), (-1, -1),
                 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("VALIGN",        (0, 0), (-1, -1),  "MIDDLE"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -2),  [cWHT, cTNT]),
            ]))
            story.append(t5)
        else:
            story.append(Paragraph("Five Cs analysis not available.", pS))
        story.append(gap(0.3))

        # ══════════════════════════════════════════════════════════════════════
        # PAGE 3 — SHAP + REASONING + SWOT
        # ══════════════════════════════════════════════════════════════════════
        story.append(PageBreak())
        story.append(shdr("SECTION 3   —   AI RISK DRIVERS (SHAP)"))
        story.append(gap(0.12))
        if pred and pred.top_shap_factors:
            WS = [4.8*cm, 2.2*cm, 2.8*cm, TW-10.2*cm]
            shap_r = [[Paragraph("Risk Factor", pWL), Paragraph("SHAP", pWC),
                       Paragraph("Direction", pWC), Paragraph("Interpretation", pWL)]]
            for f in pred.top_shap_factors:
                up = "increases" in f.direction.lower()
                shap_r.append([Paragraph(f.display_name, pN),
                               Paragraph(f"{f.shap_value:.4f}", pNC),
                               Paragraph(
                    f'<font color="{"#b71c1c" if up else "#1a6b2a"}">{"▲ Inc." if up else "▼ Dec."}</font>', pNC),
                    Paragraph("Unfavourable" if up else "Favourable", pS)])
            story.append(col_tbl(shap_r, WS))
        else:
            story.append(Paragraph("SHAP analysis not available.", pS))
        story.append(gap(0.2))

        # AI reasoning — compact
        story.append(shdr("SECTION 4   —   AI REASONING CHAIN"))
        story.append(gap(0.1))
        if result.reasoning_chain:
            key_starts = ("DECISION:", "LIMIT:", "RATE:",
                          "REASONING", "DECISIVE", "LOAN")
            rc_data = []
            for line in result.reasoning_chain.split("\n"):
                line = line.strip()
                if not line:
                    continue
                is_key = any(line.upper().startswith(k) for k in key_starts)
                rc_data.append([Paragraph(line, pB if is_key else pN)])
            if rc_data:
                rc_t = Table(rc_data, colWidths=[TW])
                rc_t.setStyle(TableStyle([
                    ("BOX", (0, 0), (-1, -1), 0.6, cNAV2), ("BACKGROUND",
                                                            (0, 0), (-1, -1), C("#f8f9ff")),
                    ("LEFTPADDING", (0, 0), (-1, -1),
                     8), ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1),
                     3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]))
                story.append(rc_t)
        else:
            story.append(Paragraph("AI reasoning not available.", pS))
        story.append(gap(0.2))

        # SWOT — compact 2x2
        story.append(shdr("SECTION 5   —   SWOT ANALYSIS"))
        story.append(gap(0.12))
        swot = getattr(result, "swot", None)
        if swot:
            def swot_q(items, title, bg):
                rows = [[Paragraph(title, PS(fontSize=7.5, fontName="Helvetica-Bold",
                                             leading=11, textColor=cWHT, alignment=TA_CENTER))]]
                for item in items:
                    rows.append([Paragraph(f"• {item}",
                                           PS(fontSize=7.5, fontName="Helvetica", leading=11, textColor=colors.black))])
                t = Table(rows, colWidths=[TW/2-0.15*cm])
                t.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), C(bg)),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [cWHT, cTNT]),
                    ("BOX", (0, 0), (-1, -1), 0.6,
                     cNAV2), ("INNERGRID", (0, 0), (-1, -1), 0.2, cMGR),
                    ("LEFTPADDING", (0, 0), (-1, -1),
                     6), ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1),
                     4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]))
                return t
            swot_grid = Table(
                [[swot_q(swot.strengths, "STRENGTHS", GRN),    swot_q(swot.weaknesses, "WEAKNESSES", AMB)],
                 [swot_q(swot.opportunities, "OPPORTUNITIES", NAV), swot_q(swot.threats, "THREATS", RED)]],
                colWidths=[TW/2, TW/2]
            )
            swot_grid.setStyle(TableStyle([
                ("LEFTPADDING", (0, 0), (-1, -1),
                 2), ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                ("TOPPADDING", (0, 0), (-1, -1),
                 2), ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]))
            story.append(swot_grid)
            if getattr(swot, "summary", ""):
                story.append(gap(0.1))
                story.append(Paragraph(swot.summary, pS))
        else:
            story.append(
                Paragraph("SWOT not available — run analysis pipeline.", pS))
        story.append(gap(0.3))

        # ══════════════════════════════════════════════════════════════════════
        # PAGE 4 — RESEARCH + TRIANGULATION + EARLY WARNINGS
        # ══════════════════════════════════════════════════════════════════════
        story.append(PageBreak())
        story.append(
            shdr("SECTION 6   —   EXTERNAL RESEARCH & MACRO TRIANGULATION"))
        story.append(gap(0.12))

        research_dict = getattr(result, "research_dict", None)
        r = result.research
        if research_dict and isinstance(research_dict, dict):
            tri = research_dict.get("triangulation", {})
            story.append(info_tbl([
                ("External Risk",    tri.get("overall_external_risk", "—")),
                ("News Risk Score",
                 f"{research_dict.get('news_risk_score','—')}/10"),
                ("Litigation",       "YES" if research_dict.get(
                    "litigation_found") else "None found"),
                ("MCA Charges",      f"{len(research_dict.get('mca_charges',[]))} charge(s)" if research_dict.get(
                    "mca_charges") else "None"),
            ], lw=3.8*cm, fs=8, pad=4))
            if tri.get("triangulation_summary"):
                story.append(gap(0.1))
                story.append(
                    Paragraph(f"<b>Synthesis:</b> {tri['triangulation_summary']}", pJ))
            # Red flags + Positives side by side
            flags = tri.get("key_red_flags", [])
            pos = tri.get("key_positives", [])
            if flags or pos:
                fl = ([Paragraph("<b>Red Flags</b>", PS(fontSize=8, fontName="Helvetica-Bold", leading=11, textColor=C(RED)))] +
                      [Paragraph(f"• {f}", PS(fontSize=7.5, fontName="Helvetica", leading=11, textColor=C(RED))) for f in flags[:3]])
                pl = ([Paragraph("<b>Positives</b>", PS(fontSize=8, fontName="Helvetica-Bold", leading=11, textColor=C(GRN)))] +
                      [Paragraph(f"• {p}", PS(fontSize=7.5, fontName="Helvetica", leading=11, textColor=C(GRN))) for p in pos[:3]])
                ft = Table([[p] for p in fl], colWidths=[TW/2-0.2*cm])
                pt = Table([[p] for p in pl], colWidths=[TW/2-0.2*cm])
                for tx in [ft, pt]:
                    tx.setStyle(TableStyle([("BOX", (0, 0), (-1, -1), 0.5, cMGR),
                                            ("LEFTPADDING", (0, 0), (-1, -1),
                                             6), ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                                            ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3)]))
                story.append(gap(0.1))
                fp_wrap = Table([[ft, pt]], colWidths=[TW/2, TW/2])
                fp_wrap.setStyle(TableStyle([("LEFTPADDING", (0, 0), (-1, -1), 0),
                                             ("RIGHTPADDING", (0, 0), (-1, -1),
                                              3), ("TOPPADDING", (0, 0), (-1, -1), 0),
                                             ("BOTTOMPADDING", (0, 0), (-1, -1), 0)]))
                story.append(fp_wrap)
        elif r:
            story.append(info_tbl([
                ("News Risk Score",   f"{r.news_risk_score}/10"),
                ("Negative Articles", str(len(r.negative_news))),
                ("Litigation",        "YES" if r.litigation_found else "None found"),
                ("MCA Charges",
                 f"{len(r.mca_charges)} charge(s)" if r.mca_charges else "None"),
            ], lw=3.8*cm, fs=8, pad=4))
            if r.research_summary:
                story.append(gap(0.1))
                story.append(Paragraph(r.research_summary, pS))
            if r.negative_news:
                story.append(gap(0.1))
                story.append(Paragraph("Adverse News:", pB))
                nr = [[Paragraph("Date", pWL), Paragraph(
                    "Headline", pWL), Paragraph("Source", pWL)]]
                for item in r.negative_news[:4]:
                    nr.append([Paragraph(str(item.date)[:10], pS), Paragraph(
                        item.title, pS), Paragraph(item.source, pS)])
                story.append(col_tbl(nr, [2*cm, 11*cm, 3*cm], hdr=RED))
        else:
            story.append(Paragraph("External research not available.", pS))
        story.append(gap(0.2))

        # Early warnings
        story.append(shdr("SECTION 7   —   EARLY WARNING SIGNALS"))
        story.append(gap(0.1))
        if pred and pred.early_warning_signals:
            freqs = ["Monthly", "Quarterly",
                     "Quarterly", "At Renewal", "Annually"]
            ew = [[Paragraph("#", pWC), Paragraph(
                "Signal", pWL), Paragraph("Frequency", pWC)]]
            for i, w in enumerate(pred.early_warning_signals):
                ew.append([Paragraph(str(i+1), pNC), Paragraph(w, pN),
                           Paragraph(freqs[i] if i < len(freqs) else "Quarterly", pNC)])
            story.append(col_tbl(ew, [0.8*cm, TW-5*cm, 3.7*cm], hdr=AMB))
        else:
            story.append(Paragraph("No early warning signals identified.", pS))
        story.append(gap(0.3))

        # ══════════════════════════════════════════════════════════════════════
        # PAGE 5 — CONDITIONS + RECOMMENDATION
        # ══════════════════════════════════════════════════════════════════════
        story.append(PageBreak())
        story.append(shdr("SECTION 8   —   CREDIT CONDITIONS & COVENANTS"))
        story.append(gap(0.1))
        conds = [
            ("Pre-Disbursement",
             "Submit audited financials + GST returns for past 3 years."),
            ("Pre-Disbursement",
             "Execute all security documents; SARFAESI-compliant first charge."),
            ("Ongoing Covenant",
             "Quarterly stock/debtor statements; maintain DSCR > 1.25x."),
            ("Ongoing Covenant",
             "Prior written approval for additional LT borrowings > Rs. 50L."),
            ("Annual Review",    "Audited Balance Sheet + P&L within 6 months of FY end."),
            ("Annual Review",    "Annual site inspection by relationship/credit manager."),
        ]
        cr = [[Paragraph("Type", pWL), Paragraph("Condition", pWL)]]
        for ct, cx in conds:
            cr.append([Paragraph(ct, PS(fontSize=7.5, fontName="Helvetica-Bold", leading=11, textColor=colors.black)),
                       Paragraph(cx, pN)])
        story.append(col_tbl(cr, [3.5*cm, TW-3.5*cm]))
        story.append(gap(0.25))

        # Recommendation
        story.append(
            shdr("SECTION 9   —   RECOMMENDATION & PROPOSED SANCTION TERMS"))
        story.append(gap(0.12))
        if pred:
            story.append(info_tbl([
                ("Borrower",          result.company_name),
                ("Sanctioned Limit",
                 f"Rs. {pred.loan_limit_inr/10000000:.2f} Cr " f"(Rs. {pred.loan_limit_inr/100000:.2f} Lakhs only)"),
                ("Rate of Interest",
                 f"{pred.interest_rate:.2f}% p.a. (Floating — MCLR linked)"),
                ("Primary Security",  "Hypothecation of current assets"),
                ("Collateral",        "As per application / valuation report"),
                ("Processing Fee",    "0.50% of sanctioned limit + GST"),
                ("Annual Review",     now.replace(
                    year=now.year+1).strftime("%d-%m-%Y")),
            ], lw=4.0*cm, pad=4))
            story.append(gap(0.15))

        # Final decision banner
        final_banner = Table([[Paragraph(f"<b>{dec_label}</b>",
                                         PS(fontSize=13, fontName="Helvetica-Bold", alignment=TA_CENTER, textColor=cWHT))]],
                             colWidths=[TW])
        final_banner.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), C(dec_bg)),
            ("TOPPADDING", (0, 0), (-1, -1),
             12), ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
        ]))
        story.append(final_banner)
        story.append(gap(0.25))

        # Four-person sign-off
        cw4 = TW/4
        sig4 = [
            [Paragraph(f"<b>{h}</b>", pWC) for h in ["Credit Analyst",
                                                     "Credit Manager", "Head of Credit", "Chief Credit Officer"]],
            [Paragraph("\n\n___________________", pNC)]*4,
            [Paragraph("Date: _____________", pS)]*4,
        ]
        st4 = Table(sig4, colWidths=[cw4]*4)
        st4.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), 0.7,
             cNAV), ("INNERGRID", (0, 0), (-1, -1), 0.3, cMGR),
            ("BACKGROUND", (0, 0), (-1, 0),
             cNAV), ("BACKGROUND", (0, 1), (-1, -1), cTNT),
            ("FONTSIZE", (0, 0), (-1, -1), 8), ("LEADING", (0, 0), (-1, -1), 11.5),
            ("TOPPADDING", (0, 0), (-1, -1),
             4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ]))
        story.append(st4)
        story.append(gap(0.3))

        # Disclaimer
        story.append(HRFlowable(width="100%", thickness=0.4, color=cMGR))
        story.append(gap(0.08))
        story.append(Paragraph(
            f"DISCLAIMER: This CAM is generated by {APP_NAME} v2.0, an AI-assisted credit "
            "scoring engine. Output is supplementary to human judgment and does not constitute "
            "a binding credit decision. All sanctions subject to applicable RBI guidelines, "
            "internal credit policy, and Credit Committee approval. Strictly confidential.",
            PS(fontSize=6.5, fontName="Helvetica", leading=10, textColor=cDGR,
               alignment=TA_JUSTIFY)))

        pdf_doc.build(story)
        print(f"PDF saved: {output_path}")
        return output_path

    # ══════════════════════════════════════════════════════════════════════════
    # DOCX  —  mirrors PDF structure, 5-6 pages
    # ══════════════════════════════════════════════════════════════════════════

    def generate_docx(self, result: CreditAppraisalResult, output_path: str) -> str:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()
        for sec in doc.sections:
            sec.top_margin = Cm(2.0)
            sec.bottom_margin = Cm(2.0)
            sec.left_margin = Cm(2.2)
            sec.right_margin = Cm(2.2)

        pred = result.risk_prediction
        ds, cs, dec_bg, dec_lt, dec_label = self._pred_strings(pred)
        now = datetime.now()
        ref = f"CAM/{now.year}/{result.company_name[:4].upper()}/{now.strftime('%m%d%H%M')}"

        def rgb(h): return RGBColor(
            int(h[1:3], 16), int(h[3:5], 16), int(h[5:7], 16))

        def shade(cell, hex_c):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), hex_c.lstrip("#"))
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            tcPr.append(shd)

        def cell_fmt(cell, text, bold=False, size=9, color=None,
                     align=WD_ALIGN_PARAGRAPH.LEFT):
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(str(text))
            run.bold = bold
            run.font.size = Pt(size)
            run.font.name = "Arial"
            if color:
                run.font.color.rgb = rgb(color)
            p.alignment = align
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

        def shdr_d(txt):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(txt)
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = "Arial"
            run.font.color.rgb = rgb(NAV2)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            lft = OxmlElement("w:left")
            lft.set(qn("w:val"), "single")
            lft.set(qn("w:sz"), "18")
            lft.set(qn("w:space"), "6")
            lft.set(qn("w:color"), NAV.lstrip("#"))
            pBdr.append(lft)
            pPr.append(pBdr)

        def info_tbl_d(rows, w0=4.0):
            t = doc.add_table(rows=0, cols=2)
            t.style = "Table Grid"
            for lbl, val in rows:
                row = t.add_row()
                c0, c1 = row.cells[0], row.cells[1]
                c0.width = Cm(w0)
                c1.width = Cm(13.5-w0)
                shade(c0, LBLU)
                cell_fmt(c0, lbl, bold=True)
                cell_fmt(c1, val)
            return t

        # Cover
        h1 = doc.add_heading(APP_NAME.upper(), level=1)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if h1.runs:
            h1.runs[0].font.color.rgb = rgb(NAV)
            h1.runs[0].font.size = Pt(14)
            h1.runs[0].font.name = "Arial"

        p_sub = doc.add_paragraph("AI-Powered Credit Intelligence Engine")
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p_sub.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = rgb(DGRY)
            r.font.name = "Arial"

        doc.add_paragraph()
        p_title = doc.add_paragraph("INVESTMENT ASSESSMENT REPORT")
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p_title.runs:
            r.bold = True
            r.font.size = Pt(17)
            r.font.color.rgb = rgb(NAV)
            r.font.name = "Arial"
        doc.add_paragraph()

        # Cover table — entity + loan details
        info_tbl_d([
            ("Company / Borrower", result.company_name),
            ("CIN",                getattr(result, "cin", "N/A")),
            ("PAN",                result.itr_data.pan if result.itr_data else "N/A"),
            ("GSTIN",              result.gst_data.gstin if result.gst_data else "N/A"),
            ("Sector",             getattr(result, "sector", "—")),
            ("Loan Type",          getattr(result, "loan_type", "—")),
            ("Amount Requested",
             f"Rs. {getattr(result,'loan_amount_cr','—')} Cr"),
            ("Tenure",
             f"{getattr(result,'loan_tenure_months','—')} months"),
            ("AI Decision",        dec_label),
            ("Risk Score",
             f"{pred.risk_score:.3f} / 1.000" if pred else "N/A"),
            ("Loan Limit",
             f"Rs. {pred.loan_limit_inr:,.0f}" if pred else "N/A"),
            ("Interest Rate",
             f"{pred.interest_rate:.2f}% p.a." if pred else "N/A"),
            ("Reference No.",      ref),
            ("Report Date",        now.strftime("%d %B %Y")),
            ("Prepared By",        CAM_AUTHOR),
            ("Classification",     "STRICTLY CONFIDENTIAL"),
        ])
        doc.add_paragraph()
        doc.add_page_break()

        # S1 Financial Analysis
        shdr_d("SECTION 1 — EXECUTIVE SUMMARY & FINANCIAL ANALYSIS")
        if pred and getattr(pred, "decisive_factor", ""):
            p = doc.add_paragraph(f"Decisive Factor: {pred.decisive_factor}")
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.name = "Arial"
        dr = self._derived_rows(result)
        info_tbl_d(dr or [("Financial Ratios", "Not available")])
        doc.add_paragraph()

        # S2 Five Cs
        shdr_d("SECTION 2 — FIVE Cs CREDIT ANALYSIS")
        if result.five_cs:
            obj5 = result.five_cs
            t5 = doc.add_table(rows=1, cols=3)
            t5.style = "Table Grid"
            for ci, h_ in enumerate(["Parameter", "Score", "Assessment"]):
                shade(t5.rows[0].cells[ci], NAV)
                cell_fmt(t5.rows[0].cells[ci], h_, bold=True, color=WHT)
            for lbl, o in [("Character", obj5.character), ("Capacity", obj5.capacity),
                           ("Capital", obj5.capital), ("Collateral", obj5.collateral),
                           ("Conditions", obj5.conditions)]:
                row = t5.add_row().cells
                cell_fmt(row[0], lbl, bold=True)
                cell_fmt(row[1], f"{o.score}/10",
                         align=WD_ALIGN_PARAGRAPH.CENTER)
                cell_fmt(row[2], o.summary)
            ov = t5.add_row().cells
            for c in ov:
                shade(c, LBLU)
            cell_fmt(ov[0], "OVERALL", bold=True)
            cell_fmt(ov[1], f"{obj5.overall_score}/10",
                     bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            cell_fmt(ov[2], "Weighted average", bold=True)
        doc.add_paragraph()

        # S3 SHAP
        shdr_d("SECTION 3 — AI RISK DRIVERS (SHAP)")
        if pred and pred.top_shap_factors:
            t4 = doc.add_table(rows=1, cols=3)
            t4.style = "Table Grid"
            for ci, h_ in enumerate(["Factor", "SHAP Impact", "Direction"]):
                shade(t4.rows[0].cells[ci], NAV)
                cell_fmt(t4.rows[0].cells[ci], h_, bold=True, color=WHT)
            for f in pred.top_shap_factors:
                row = t4.add_row().cells
                cell_fmt(row[0], f.display_name)
                cell_fmt(row[1], f"{f.shap_value:.4f}",
                         align=WD_ALIGN_PARAGRAPH.CENTER)
                cell_fmt(row[2], f.direction)
        doc.add_paragraph()

        # S4 SWOT
        shdr_d("SECTION 4 — SWOT ANALYSIS")
        swot = getattr(result, "swot", None)
        if swot:
            for cat, items in [("STRENGTHS", swot.strengths), ("WEAKNESSES", swot.weaknesses),
                               ("OPPORTUNITIES", swot.opportunities), ("THREATS", swot.threats)]:
                p = doc.add_paragraph(cat)
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)
                    r.font.name = "Arial"
                for item in items:
                    bi = doc.add_paragraph(f"  • {item}")
                    for r in bi.runs:
                        r.font.size = Pt(8.5)
                        r.font.name = "Arial"
                    bi.paragraph_format.space_after = Pt(1)
            if getattr(swot, "summary", ""):
                doc.add_paragraph(swot.summary)
        else:
            doc.add_paragraph("SWOT analysis not available.")
        doc.add_paragraph()

        # S5 Research
        shdr_d("SECTION 5 — EXTERNAL RESEARCH")
        r = result.research
        if r:
            info_tbl_d([
                ("News Risk Score",  f"{r.news_risk_score}/10"),
                ("Negative News",   str(len(r.negative_news))),
                ("Litigation",      "YES" if r.litigation_found else "None"),
                ("MCA Charges",
                 f"{len(r.mca_charges)} charge(s)" if r.mca_charges else "None"),
            ])
            if r.research_summary:
                doc.add_paragraph(r.research_summary)
        else:
            doc.add_paragraph("Research not available.")
        doc.add_paragraph()

        # S6 Early Warnings
        shdr_d("SECTION 6 — EARLY WARNING SIGNALS")
        if pred and pred.early_warning_signals:
            t5w = doc.add_table(rows=1, cols=2)
            t5w.style = "Table Grid"
            for ci, h_ in enumerate(["Signal", "Frequency"]):
                shade(t5w.rows[0].cells[ci], AMB)
                cell_fmt(t5w.rows[0].cells[ci], h_, bold=True, color=WHT)
            freqs = ["Monthly", "Quarterly",
                     "Quarterly", "At Renewal", "Annually"]
            for i, w in enumerate(pred.early_warning_signals):
                row = t5w.add_row().cells
                cell_fmt(row[0], w)
                cell_fmt(row[1], freqs[i] if i < len(freqs) else "Quarterly",
                         align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()

        # S7 Recommendation
        shdr_d("SECTION 7 — RECOMMENDATION")
        if pred:
            info_tbl_d([
                ("Borrower",         result.company_name),
                ("Sanctioned Limit", f"Rs. {pred.loan_limit_inr:,.0f}"),
                ("Interest Rate",    f"{pred.interest_rate:.2f}% p.a."),
                ("Primary Security", "Hypothecation of current assets"),
                ("Processing Fee",   "0.50% + GST"),
            ])
            p_dec = doc.add_paragraph()
            r_dec = p_dec.add_run(f"DECISION: {dec_label}")
            r_dec.bold = True
            r_dec.font.size = Pt(12)
            r_dec.font.color.rgb = rgb(dec_bg)
            r_dec.font.name = "Arial"
            if pred.decisive_factor:
                p_df = doc.add_paragraph(
                    f"Decisive Factor: {pred.decisive_factor}")
                for r2 in p_df.runs:
                    r2.font.size = Pt(9)
                    r2.font.name = "Arial"

        # Footer
        footer = doc.sections[0].footer
        fp = footer.paragraphs[0]
        fp.text = (f"Ref: {ref}  |  {APP_NAME}  |  "
                   f"STRICTLY CONFIDENTIAL  |  {now.strftime('%d %b %Y')}  |  {CAM_AUTHOR}")
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if fp.runs:
            fp.runs[0].font.size = Pt(7)
            fp.runs[0].font.color.rgb = rgb(DGRY)
            fp.runs[0].font.name = "Arial"

        doc.save(output_path)
        print(f"DOCX saved: {output_path}")
        return output_path
