# src/classifier.py
import json
import re
from pathlib import Path
from typing import Optional
from src.schemas import DocumentClassification

CLASSIFICATION_SIGNALS = {
    "ANNUAL_REPORT": [
        "profit and loss", "balance sheet", "cash flow statement",
        "independent auditor", "board's report", "revenue from operations",
        "total assets", "total liabilities", "earnings per share",
        "statement of profit", "notes to accounts", "schedule iii"
    ],
    "ALM": [
        "asset liability", "maturity profile", "bucket", "outflows",
        "inflows", "1-30 days", "31-60", "over 5 years",
        "structural liquidity", "dynamic liquidity", "alm committee",
        "rate sensitive", "interest rate sensitivity", "gap analysis"
    ],
    "SHAREHOLDING_PATTERN": [
        "promoter", "public shareholding", "pledge", "encumbered",
        "demat", "physical shares", "institutional investors",
        "foreign portfolio", "mutual fund holding", "shareholding pattern",
        "promoter group", "non-promoter", "depository receipts"
    ],
    "BORROWING_PROFILE": [
        "borrowings", "debentures", "ncd", "term loan lenders",
        "debt funding", "credit rating", "subordinated debt",
        "external commercial borrowing", "lender", "sanction limit",
        "outstanding amount", "borrowing mix", "secured borrowings",
        "unsecured borrowings", "interest coverage"
    ],
    "PORTFOLIO_PERFORMANCE": [
        "aum", "assets under management", "disbursement", "gnpa",
        "nnpa", "collection efficiency", "portfolio at risk",
        "par 30", "par 60", "loan book", "npa", "write-off",
        "yield on portfolio", "cost of funds", "net interest margin",
        "credit cost", "stage 1", "stage 2", "stage 3"
    ]
}

LABEL_MAP = {
    "ANNUAL_REPORT":         "Annual Report (P&L / Balance Sheet / Cashflow)",
    "ALM":                   "ALM (Asset-Liability Management)",
    "SHAREHOLDING_PATTERN":  "Shareholding Pattern",
    "BORROWING_PROFILE":     "Borrowing Profile",
    "PORTFOLIO_PERFORMANCE": "Portfolio / Performance Data",
    "UNKNOWN":               "Unknown — needs manual classification"
}

ALL_DOC_TYPES = list(LABEL_MAP.keys())

# ── Per-doc-type extraction keywords ──────────────────────────────────────────
# More specific than FINANCIAL_KEYWORDS — used to score pages per doc type
DOC_TYPE_PAGE_KEYWORDS = {
    "ANNUAL_REPORT": [
        "profit and loss", "balance sheet", "cash flow",
        "revenue from operations", "total income", "total assets",
        "net worth", "borrowings", "pat", "ebitda",
        "interest expense", "finance cost", "total equity",
        "net profit", "earnings per share", "notes to financial",
        "statement of profit", "other comprehensive income",
        "deferred tax", "capital work in progress"
    ],
    "ALM": [
        "maturity", "asset liability", "bucket", "1-30", "31-60",
        "outflow", "inflow", "liquidity gap", "cumulative gap",
        "rate sensitive", "gap analysis", "structural liquidity"
    ],
    "SHAREHOLDING_PATTERN": [
        "promoter", "public shareholding", "pledge", "demat",
        "institutional", "foreign portfolio", "mutual fund",
        "non-promoter", "encumbered", "category"
    ],
    "BORROWING_PROFILE": [
        "borrowings", "ncd", "debentures", "term loan", "credit rating",
        "lender", "outstanding", "secured", "unsecured", "ecb",
        "subordinated", "interest coverage", "debt"
    ],
    "PORTFOLIO_PERFORMANCE": [
        "aum", "disbursement", "gnpa", "nnpa", "collection efficiency",
        "par 30", "par 60", "loan book", "npa", "yield",
        "cost of funds", "nim", "stage 1", "stage 2", "stage 3",
        "credit cost", "roe", "roa", "capital adequacy"
    ],
}

# Generic financial keywords — used when doc type is unknown
FINANCIAL_KEYWORDS = [
    "profit and loss", "balance sheet", "cash flow",
    "revenue from operations", "total income", "total assets",
    "net worth", "borrowings", "aum", "gnpa", "disbursement",
    "promoter", "shareholding", "maturity", "asset liability",
    "pat", "ebitda", "interest expense", "finance cost",
    "total equity", "net profit", "earnings per share",
    "npa", "collection efficiency", "loan book", "ncd",
    "credit rating", "lender", "stage 1", "stage 2", "stage 3"
]


def extract_preview_text(file_path: str, max_chars: int = 3000) -> str:
    """Fast text preview for classification — first N pages only."""
    suffix = Path(file_path).suffix.lower()
    try:
        if suffix == ".pdf":
            import fitz
            doc = fitz.open(file_path)
            text = ""
            for page in doc[:15]:
                text += page.get_text()
                if len(text) >= max_chars:
                    break
            doc.close()
            return text[:max_chars]

        elif suffix in (".xlsx", ".xls"):
            import openpyxl
            wb = openpyxl.load_workbook(
                file_path, read_only=True, data_only=True)
            text = ""
            for sheet in wb.sheetnames[:3]:
                ws = wb[sheet]
                for row in ws.iter_rows(max_row=30, values_only=True):
                    text += " ".join(str(c) for c in row if c) + "\n"
            wb.close()
            return text[:max_chars]

        elif suffix == ".docx":
            from docx import Document
            doc = Document(file_path)
            text = "\n".join(p.text for p in doc.paragraphs[:50])
            return text[:max_chars]

    except Exception as e:
        return f"preview_error: {str(e)}"
    return ""


def extract_financial_text(
    file_path: str,
    # FIX 1: 12k → 40k (covers real financial tables)
    max_chars: int = 40000,
    doc_type_hint: str = None,     # FIX 2: use doc type to score pages more precisely
    max_pages: int = 40            # FIX 3: 15 → 40 financial pages
) -> str:
    """
    Targeted extraction for financial data — scans ALL pages, keeps only
    pages containing financial keywords relevant to the document type.

    Improvements over v1:
    - max_chars 40k instead of 12k — covers full P&L + Balance Sheet + Cashflow
    - max_pages 40 instead of 15 — annual reports have many financial pages
    - doc_type_hint scoring — uses type-specific keywords for better page selection
    - Deduplicates near-identical pages (common in annual reports)
    """
    suffix = Path(file_path).suffix.lower()

    # For Excel/Word — extract all sheets, no page limit needed
    if suffix in (".xlsx", ".xls"):
        return _extract_excel_full(file_path, max_chars)
    if suffix == ".docx":
        return _extract_docx_full(file_path, max_chars)
    if suffix != ".pdf":
        return extract_preview_text(file_path, max_chars)

    try:
        import fitz
        doc = fitz.open(file_path)
        total_pages = len(doc)

        # Choose scoring keywords based on doc type hint
        if doc_type_hint and doc_type_hint in DOC_TYPE_PAGE_KEYWORDS:
            scoring_kw = DOC_TYPE_PAGE_KEYWORDS[doc_type_hint]
        else:
            scoring_kw = FINANCIAL_KEYWORDS

        financial_pages = []
        cover_texts = []
        seen_snippets = set()   # deduplication

        for i, page in enumerate(doc):
            text = page.get_text("text")
            if not text.strip():
                continue

            text_lower = text.lower()
            hits = sum(1 for kw in scoring_kw if kw in text_lower)

            # Always include first 5 pages for cover/context
            if i < 5:
                cover_texts.append(text)
                continue

            if hits >= 2:
                # Dedup — skip if this page's first 100 chars already seen
                snippet = text_lower[:100].strip()
                if snippet in seen_snippets:
                    continue
                seen_snippets.add(snippet)
                financial_pages.append((i, hits, text))

        doc.close()

        # Sort by relevance score
        financial_pages.sort(key=lambda x: x[1], reverse=True)

        print(f"[EXTRACT_TEXT] {total_pages} pages total, "
              f"{len(financial_pages)} financial pages found, "
              f"using top {min(max_pages, len(financial_pages))}")

        # Build output: cover + top financial pages
        combined = "\n".join(cover_texts)
        combined += "\n\n--- FINANCIAL SECTIONS ---\n\n"
        for _, hits, text in financial_pages[:max_pages]:
            combined += text + "\n\n"
            if len(combined) >= max_chars:
                break

        return combined[:max_chars]

    except Exception as e:
        print(f"[EXTRACT_TEXT] Error: {e}, falling back to preview")
        return extract_preview_text(file_path, max_chars)


def _extract_excel_full(file_path: str, max_chars: int = 40000) -> str:
    """Extract ALL sheets from Excel — no row limit."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        text = ""
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            text += f"\n--- Sheet: {sheet} ---\n"
            for row in ws.iter_rows(values_only=True):
                row_vals = [str(c)
                            for c in row if c is not None and str(c).strip()]
                if row_vals:
                    text += "  ".join(row_vals) + "\n"
                if len(text) >= max_chars:
                    break
            if len(text) >= max_chars:
                break
        wb.close()
        return text[:max_chars]
    except Exception as e:
        return f"excel_error: {e}"


def _extract_docx_full(file_path: str, max_chars: int = 40000) -> str:
    """Extract all paragraphs from DOCX."""
    try:
        from docx import Document
        doc = Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return text[:max_chars]
    except Exception as e:
        return f"docx_error: {e}"


def classify_by_keywords(text: str) -> tuple[str, float, dict]:
    """Fast keyword pre-classifier. Returns (doc_type, confidence, score_map)."""
    text_lower = text.lower()
    scores = {
        doc_type: sum(1 for kw in keywords if kw in text_lower)
        for doc_type, keywords in CLASSIFICATION_SIGNALS.items()
    }
    best = max(scores, key=scores.get)
    total_hits = sum(scores.values())

    if scores[best] < 2 or total_hits == 0:
        return "UNKNOWN", 0.15, scores

    confidence = round(min(scores[best] / total_hits, 0.95), 2)
    return best, confidence, scores


def classify_with_llm(
    filename: str,
    text_preview: str,
    kw_hint: str
) -> DocumentClassification:
    """LLM-based classifier with keyword hint. Falls back gracefully."""
    from config import get_groq_client
    client = get_groq_client()

    # FIX 4: send 1500 chars to LLM — enough signal, not too expensive
    prompt = f"""You are a financial document classifier for an NBFC credit underwriting system.

Classify the document into EXACTLY ONE of:
- ANNUAL_REPORT    : P&L, Balance Sheet, Cashflow, Auditor Report, Board Report
- ALM              : Asset-Liability table with maturity buckets / time bands
- SHAREHOLDING_PATTERN : Promoter %, public holding %, pledge data
- BORROWING_PROFILE : Lender list, NCD details, borrowing mix, credit ratings
- PORTFOLIO_PERFORMANCE : AUM, disbursements, NPA%, collection efficiency
- UNKNOWN          : Cannot be determined

Filename: {filename}
Keyword pre-classification hint: {kw_hint}

Document preview (first 1500 chars):
{text_preview[:1500]}

Respond ONLY with valid JSON — no markdown, no explanation outside the JSON:
{{
  "doc_type": "<type>",
  "confidence": <0.0-1.0>,
  "reasoning": "<one sentence>",
  "key_signals": ["<signal1>", "<signal2>", "<signal3>"]
}}"""

    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=200
        )
        raw = response.choices[0].message.content.strip()
        raw = re.sub(r"```(?:json)?|```", "", raw).strip()
        data = json.loads(raw)

        doc_type = data.get("doc_type", kw_hint)
        if doc_type not in ALL_DOC_TYPES:
            doc_type = "UNKNOWN"

        return DocumentClassification(
            doc_type=doc_type,
            doc_type_label=LABEL_MAP[doc_type],
            confidence=float(data.get("confidence", 0.5)),
            reasoning=data.get("reasoning", "LLM classification"),
            key_signals=data.get("key_signals", [])
        )

    except Exception as e:
        kw_type, kw_conf, _ = classify_by_keywords(text_preview)
        return DocumentClassification(
            doc_type=kw_type,
            doc_type_label=LABEL_MAP.get(kw_type, "Unknown"),
            confidence=kw_conf,
            reasoning=f"LLM unavailable — keyword fallback. ({str(e)[:60]})",
            key_signals=[]
        )


def classify_document(filename: str, parsed_text: str) -> DocumentClassification:
    """
    Main entry point. Keyword pass first — if confidence ≥ 0.75, skip LLM.
    Otherwise calls LLM with keyword hint.
    """
    kw_type, kw_conf, kw_scores = classify_by_keywords(parsed_text)

    if kw_conf >= 0.75:
        matched_signals = [
            kw for kw in CLASSIFICATION_SIGNALS.get(kw_type, [])
            if kw in parsed_text.lower()
        ][:4]
        return DocumentClassification(
            doc_type=kw_type,
            doc_type_label=LABEL_MAP[kw_type],
            confidence=kw_conf,
            reasoning=f"High-confidence keyword match ({kw_conf:.0%}).",
            key_signals=matched_signals
        )

    return classify_with_llm(filename, parsed_text, kw_type)
